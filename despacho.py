import streamlit as st
import pandas as pd
import numpy as np
from datetime import datetime, timedelta
import urllib.parse
import requests
import io
import re
import plotly.express as px
import zipfile
from docx import Document
from docx.shared import Inches
import plotly.graph_objects as go
from plotly.subplots import make_subplots

# Código CSS para ocultar elementos de la interfaz
hide_github_style = """
    <style>
    /* Ocultar el botón de GitHub/Fork en la esquina superior derecha */
    #GithubIcon {visibility: hidden;}
    
    /* Ocultar el menú de hamburguesa estándar de Streamlit */
    #MainMenu {visibility: hidden;}
    
    /* Ocultar la barra de encabezado superior por completo */
    header {visibility: hidden;}
    
    /* Ocultar el pie de página "Made with Streamlit" */
    footer {visibility: hidden;}
    </style>
"""

# Inyectar CSS en la aplicación
st.markdown(
    """
    <style>
    .css-1jc7ptx, .e1ewe7hr3, .viewerBadge_container__1QSob,
    .styles_viewerBadge__1yB5_, .viewerBadge_link__1S137,
    .viewerBadge_text__1JaDK {
        display: none;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# --- 1. CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="Supervisión Despacho - SEIN", layout="wide", initial_sidebar_state="expanded")
st.title("⚡ Dashboard de Supervisión - Despacho Ejecutado del SEIN ")
st.markdown("Supervisión del Despacho, Interconexiones y Seguridad Operativa")

MESES = {
    1: "Enero", 2: "Febrero", 3: "Marzo", 4: "Abril",
    5: "Mayo", 6: "Junio", 7: "Julio", 8: "Agosto",
    9: "Setiembre", 10: "Octubre", 11: "Noviembre", 12: "Diciembre"
}

# --- CARGAR MATRIZ DE CENTRALES DEL SEIN ---
@st.cache_data
def cargar_centrales_sein():
    try:
        # Se ancla la extracción exactamente a las columnas A, B, C, D, E, G, H, I 
        df_centrales = pd.read_excel('CetralesSEIN.xlsx', sheet_name=0, header=None, usecols=[0, 1, 2, 3, 4, 6, 7, 8])
        
        # Saltar fila de encabezado (fila 0)
        df_centrales_limpio = df_centrales.iloc[1:].copy()
        
        df_centrales_limpio.columns = [
            'CODIGO', 'CENTRAL', 'CENTRAL_CALIFICACION', 'EMPRESA_DESPACHO', 
            'AREA_OPERATIVA', 'TIPO_INTEGRANTE', 'TIPO_GENERACION', 'REQUERIMIENTO_ESPECIAL'
        ]
        
        # Limpiar espacios y celdas nulas
        for col in df_centrales_limpio.columns:
            df_centrales_limpio[col] = df_centrales_limpio[col].apply(lambda x: str(x).strip() if pd.notna(x) and str(x) != 'nan' else '')
        
        # Filtrar solo registros donde el nombre de la central sea válido
        df_centrales_limpio = df_centrales_limpio[df_centrales_limpio['CENTRAL'] != ''].copy()
        
        return df_centrales_limpio
    except Exception as e:
        st.error(f"Error cargando CetralesSEIN.xlsx: {e}")
        return pd.DataFrame()

# Cargar al iniciar
df_matriz_centrales = cargar_centrales_sein()

# Crear un diccionario global de tipos de recursos basados en la Columna H
if not df_matriz_centrales.empty:
    dict_recursos_maestro = dict(zip(
        df_matriz_centrales['CENTRAL'].str.strip().str.upper(), 
        df_matriz_centrales['TIPO_GENERACION'].str.strip().str.upper()
    ))
else:
    dict_recursos_maestro = {}

# --- 2. FUNCIONES DE EXTRACCIÓN Y LIMPIEZA (ETL) ---
def generar_urls_coes(fecha):
    año = fecha.strftime("%Y")
    mes_num = fecha.strftime("%m")
    dia = fecha.strftime("%d")
    mes_titulo = MESES[fecha.month]
    fecha_str = fecha.strftime("%d%m")
    
    path_nuevo = f"Post Operación/Reportes/IEOD/{año}/{mes_num}_{mes_titulo}/{dia}/AnexoA_{fecha_str}.xlsx"
    path_legacy = f"Post Operación/Reportes/IEOD/{año}/{mes_num}_{mes_titulo}/{dia}/Anexo1_Resumen_{fecha_str}.xlsx"
    
    return [
        (f"https://www.coes.org.pe/portal/browser/download?url={urllib.parse.quote(path_nuevo)}", "AnexoA"),
        (f"https://www.coes.org.pe/portal/browser/download?url={urllib.parse.quote(path_legacy)}", "Anexo1")
    ]

def parse_dates_coes(val, base_date):
    if pd.isna(val): return pd.NaT
    if isinstance(val, (datetime, pd.Timestamp)): return pd.to_datetime(val)
    try:
        parsed = pd.to_datetime(str(val))
        if parsed.year == datetime.now().year and parsed.month == datetime.now().month and parsed.day == datetime.now().day:
            return pd.Timestamp(datetime.combine(base_date.date(), parsed.time()))
        return parsed
    except:
        return pd.NaT

@st.cache_data(show_spinner=False)
def extraer_datos_despacho(fecha, dict_recursos):
    urls = generar_urls_coes(fecha)
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
    }
    
    errores_tecnicos = []
    df_melt = pd.DataFrame()
    df_inter_melt = pd.DataFrame()
    df_seguridad_dia = pd.DataFrame()
    df_demanda_dia = pd.DataFrame()
    
    for url, tipo_anexo in urls:
        try:
            res = requests.get(url, headers=headers, timeout=20)
            if res.status_code == 200:
                try:
                    archivo_excel = io.BytesIO(res.content)
                    xls = pd.ExcelFile(archivo_excel, engine='openpyxl')
                except Exception as e:
                    errores_tecnicos.append(f"{tipo_anexo}: El archivo en la URL no es un Excel válido ({type(e).__name__})")
                    continue
                    
                hojas_limpias = {h.strip().upper(): h for h in xls.sheet_names}
                
                # ==========================================
                # EXTRACCIÓN HOJA 1: DESPACHO_EJECUTADO
                # ==========================================
                if "DESPACHO_EJECUTADO" in hojas_limpias:
                    nombre_real = hojas_limpias["DESPACHO_EJECUTADO"]
                    df_raw = pd.read_excel(xls, sheet_name=nombre_real, header=None)
                    
                    if tipo_anexo == "AnexoA":
                        zonas_raw = df_raw.iloc[6, 2:].values
                        tipos_raw = df_raw.iloc[7, 2:].values
                        empresas_raw = df_raw.iloc[8, 2:].values
                        plantas_raw = df_raw.iloc[9, 2:].values
                        data_raw = df_raw.iloc[10:58, 2:].values
                    else:
                        plantas_raw = df_raw.iloc[9, 1:].values
                        zonas_raw = df_raw.iloc[6, 1:1 + len(plantas_raw)].values
                        tipos_raw = ["N/A"] * len(plantas_raw)
                        empresas_raw = ["N/A"] * len(plantas_raw)
                        data_raw = df_raw.iloc[10:58, 1:1 + len(plantas_raw)].values
                    
                    idx_validos, nombres_plantas, dict_metadatos = [], [], {}
                    
                    for i, p in enumerate(plantas_raw):
                        if pd.notna(p):
                            nombre_base_crudo = str(p).strip().upper()
                            
                            if nombre_base_crudo != '' and 'MW' not in nombre_base_crudo:
                                tipo_excel = str(tipos_raw[i]).strip().upper() if pd.notna(tipos_raw[i]) else "N/A"
                                
                                # Buscar el tipo exacto en la Columna H de la Matriz SEIN
                                tipo_real = dict_recursos.get(nombre_base_crudo, tipo_excel)
                                
                                # Solo aplicamos fallback si el maestro no tiene info detallada
                                if tipo_real in ["N/A", "TERMOELÉCTRICA", "TERMOELECTRICA"]:
                                    centrales_biomasa = ["MAPLE ETANOL", "SAN JACINTO", "AGROLMOS", "CAÑA BRAVA", "CASA GRANDE"]
                                    centrales_diesel = ["TUMBES MAK 1", "TUMBES MAK 2", "RF GENERACION TALARA", "RECKA", "RF ETEN TG1", "RF ETEN TG2"]
                                    centrales_gas = ["MALACAS1 TG6", "MALACAS2 TG4", "REFINERIA TALARA TV1", "REFINERÍA TALARA TV2"]
                                    
                                    if nombre_base_crudo in centrales_biomasa:
                                        tipo_real = "BIOMASA"
                                    elif nombre_base_crudo in centrales_diesel:
                                        tipo_real = "DIESEL/RESIDUAL"
                                    elif nombre_base_crudo in centrales_gas:
                                        tipo_real = "GAS"
                                
                                # Preservar el agrupamiento exacto del Excel
                                tipo_agrupado = tipo_real

                                # Determinar abreviatura para sufijo de central
                                if "BIOMASA" in tipo_real: abrev = "BIO"
                                elif "GAS" in tipo_real: abrev = "GAS"
                                elif "DIESEL" in tipo_real or "RESIDUAL" in tipo_real: abrev = "DIE"
                                elif "HIDRO" in tipo_real: abrev = "HID"
                                elif "SOLAR" in tipo_real: abrev = "SOL"
                                elif "EOL" in tipo_real or "EÓL" in tipo_real: abrev = "EOL"
                                else: abrev = tipo_real[:3]

                                nombre_central = f"{nombre_base_crudo} ({abrev})"
                                
                                idx_validos.append(i)
                                nombres_plantas.append(nombre_central)
                                dict_metadatos[nombre_central] = {
                                    'ZONA': str(zonas_raw[i]).strip().upper() if pd.notna(zonas_raw[i]) else "N/A",
                                    'TIPO_CENTRAL': tipo_agrupado,
                                    'EMPRESA': str(empresas_raw[i]).strip().upper() if pd.notna(empresas_raw[i]) else "N/A"
                                }
                    
                    datos_limpios = data_raw[:, idx_validos]
                    df_dia_des = pd.DataFrame(datos_limpios, columns=nombres_plantas)
                    fechas_horas = [fecha + timedelta(minutes=30 * (i + 1)) for i in range(48)]
                    df_dia_des['FECHA_HORA'] = fechas_horas
                    
                    df_melt = df_dia_des.melt(id_vars=['FECHA_HORA'], var_name='CENTRAL', value_name='DESPACHO_MW')
                    df_melt['DESPACHO_MW'] = pd.to_numeric(df_melt['DESPACHO_MW'], errors='coerce').fillna(0)
                    df_melt['ZONA'] = df_melt['CENTRAL'].map(lambda x: dict_metadatos[x]['ZONA'])
                    df_melt['TIPO_CENTRAL'] = df_melt['CENTRAL'].map(lambda x: dict_metadatos[x]['TIPO_CENTRAL'])
                    df_melt['EMPRESA'] = df_melt['CENTRAL'].map(lambda x: dict_metadatos[x]['EMPRESA'])
                    df_melt = df_melt.groupby(['FECHA_HORA', 'CENTRAL', 'ZONA', 'TIPO_CENTRAL', 'EMPRESA'], as_index=False)['DESPACHO_MW'].sum()
                
                # ==========================================
                # EXTRACCIÓN HOJA: DEMANDA_AREAS
                # ==========================================
                if "DEMANDA_AREAS" in hojas_limpias:
                    try:
                        hoja_dem = hojas_limpias["DEMANDA_AREAS"]
                        df_raw_dem = pd.read_excel(xls, sheet_name=hoja_dem, header=None)
                        data_dem = df_raw_dem.iloc[7:55, 4].values
                        df_demanda_dia = pd.DataFrame({
                            'FECHA_HORA': [fecha + timedelta(minutes=30 * (i + 1)) for i in range(48)],
                            'DEMANDA_MW': data_dem
                        })
                        df_demanda_dia['DEMANDA_MW'] = pd.to_numeric(df_demanda_dia['DEMANDA_MW'], errors='coerce').fillna(0)
                    except Exception as e_dem:
                        pass
                        
                # ==========================================
                # EXTRACCIÓN HOJA: ENLACES (ANTES INTERCONEXIONES)
                # ==========================================
                if "INTERCONEXIONES" in hojas_limpias:
                    try:
                        hoja_inter = hojas_limpias["INTERCONEXIONES"]
                        df_raw_inter = pd.read_excel(xls, sheet_name=hoja_inter, header=None)
                        
                        # Extraer Enlace Centro-Norte (Columnas C a G -> Índices 2 a 6)
                        nom_cn = df_raw_inter.iloc[6, 2:7].values
                        nom_cn = [str(x).strip() if pd.notna(x) else f"LÍNEA_CN_{idx+1}" for idx, x in enumerate(nom_cn)]
                        data_cn = df_raw_inter.iloc[7:55, 2:7].values
                        df_cn = pd.DataFrame(data_cn, columns=nom_cn)
                        df_cn['ENLACE'] = 'CENTRO-NORTE'
                        
                        # Extraer Enlace Centro-Sur (Columnas I a L -> Índices 8 a 11)
                        nom_cs = df_raw_inter.iloc[6, 8:12].values
                        nom_cs = [str(x).strip() if pd.notna(x) else f"LÍNEA_CS_{idx+1}" for idx, x in enumerate(nom_cs)]
                        data_cs = df_raw_inter.iloc[7:55, 8:12].values
                        df_cs = pd.DataFrame(data_cs, columns=nom_cs)
                        df_cs['ENLACE'] = 'CENTRO-SUR'
                        
                        # Generar horas y consolidar ambas tablas
                        fechas_horas = [fecha + timedelta(minutes=30 * (i + 1)) for i in range(48)]
                        df_cn['FECHA_HORA'] = fechas_horas
                        df_cs['FECHA_HORA'] = fechas_horas
                        
                        df_melt_cn = df_cn.melt(id_vars=['FECHA_HORA', 'ENLACE'], var_name='LINEA_TRANSMISION', value_name='FLUJO_MW')
                        df_melt_cs = df_cs.melt(id_vars=['FECHA_HORA', 'ENLACE'], var_name='LINEA_TRANSMISION', value_name='FLUJO_MW')
                        
                        df_inter_melt = pd.concat([df_melt_cn, df_melt_cs], ignore_index=True)
                        df_inter_melt['FLUJO_MW'] = pd.to_numeric(df_inter_melt['FLUJO_MW'], errors='coerce').fillna(0)
                    except Exception as e_inter:
                        pass
                # ==========================================
                # EXTRACCIÓN HOJA: DEMANDA_AREAS
                # ==========================================
                if "DEMANDA_AREAS" in hojas_limpias:
                    try:
                        hoja_dem = hojas_limpias["DEMANDA_AREAS"]
                        df_raw_dem = pd.read_excel(xls, sheet_name=hoja_dem, header=None)
                        
                        # Extraemos SEIN, NORTE, CENTRO y SUR (Omitiendo explícitamente a SOBREANDES)
                        data_dem = df_raw_dem.iloc[7:55, 3:7].values
                        df_dem_area = pd.DataFrame(data_dem, columns=["SEIN", "NORTE", "CENTRO", "SUR"])
                        df_dem_area['FECHA_HORA'] = [fecha + timedelta(minutes=30 * (i + 1)) for i in range(48)]
                        
                        df_demanda_dia = df_dem_area.melt(id_vars=['FECHA_HORA'], var_name='ÁREA', value_name='DEMANDA_MW')
                        df_demanda_dia['DEMANDA_MW'] = pd.to_numeric(df_demanda_dia['DEMANDA_MW'], errors='coerce').fillna(0)
                    except Exception as e_dem:
                        pass


                # ==========================================
                # EXTRACCIÓN HOJA: CALIFICA_OPE_UG (CALIFICACIÓN DE OPERACIONES)
                # ==========================================
                if "CALIFICA_OPE_UG" in hojas_limpias:
                    try:
                        hoja_cal = hojas_limpias["CALIFICA_OPE_UG"]
                        df_raw_cal = pd.read_excel(xls, sheet_name=hoja_cal, header=None)
                        df_califica = df_raw_cal.iloc[6:, [1, 2, 3, 4, 5, 6, 9]].copy()
                        df_califica.columns = ['EMPRESA', 'CENTRAL', 'GRUPO', 'MODO_OPERACION', 'INICIO', 'FIN', 'TIPO_OPERACION']
                        df_califica = df_califica.dropna(how='all')
                        df_califica['TIPO_OPERACION'] = df_califica['TIPO_OPERACION'].astype(str).str.upper().str.strip()
                        
                        df_seguridad_dia = df_califica.copy()
                        if not df_seguridad_dia.empty:
                            df_seguridad_dia['INICIO'] = df_seguridad_dia['INICIO'].apply(lambda x: parse_dates_coes(x, fecha))
                            df_seguridad_dia['FIN'] = df_seguridad_dia['FIN'].apply(lambda x: parse_dates_coes(x, fecha))
                            df_seguridad_dia['FECHA_REPORTE'] = fecha.date()
                    except Exception as e_cal:
                        pass

                return df_melt, df_inter_melt, df_seguridad_dia, df_demanda_dia, None
                
            else:
                errores_tecnicos.append(f"{tipo_anexo}: HTTP {res.status_code}")
                
        except Exception as e:
            errores_tecnicos.append(f"Error nativo en {tipo_anexo}: {type(e).__name__} - {str(e)}")
            continue
            
    motivo_error = " | ".join(errores_tecnicos) if errores_tecnicos else "Falla desconocida."
    return pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), pd.DataFrame(), f"[{fecha.strftime('%d/%m/%Y')}] ❌ Motivo: {motivo_error}"

def procesar_rango_fechas(start_date, end_date, progress_bar, status_text, dict_recursos):
    fechas = pd.date_range(start_date, end_date)
    total_dias = len(fechas)
    lista_dfs, lista_inter, lista_seguridad, lista_demanda, alertas = [], [], [], [], []
    
    for i, f in enumerate(fechas):
        status_text.markdown(f"**⏳ Sincronizando datos de Despacho (COES):** {f.strftime('%d/%m/%Y')} *(Día {i+1} de {total_dias})*")
        df_dia, df_inter_dia, df_seguridad_dia, df_demanda_dia, error = extraer_datos_despacho(f, dict_recursos)
        
        if not df_dia.empty: lista_dfs.append(df_dia)
        if not df_inter_dia.empty: lista_inter.append(df_inter_dia)
        if not df_seguridad_dia.empty: lista_seguridad.append(df_seguridad_dia)
        if not df_demanda_dia.empty: lista_demanda.append(df_demanda_dia)
        if error: alertas.append(error)
            
        progress_bar.progress((i + 1) / total_dias)
            
    df_final = pd.concat(lista_dfs, ignore_index=True) if lista_dfs else pd.DataFrame()
    df_inter_final = pd.concat(lista_inter, ignore_index=True) if lista_inter else pd.DataFrame()
    df_seguridad_final = pd.concat(lista_seguridad, ignore_index=True) if lista_seguridad else pd.DataFrame()
    df_demanda_final = pd.concat(lista_demanda, ignore_index=True) if lista_demanda else pd.DataFrame()
    
    return df_final, df_inter_final, df_seguridad_final, df_demanda_final, alertas
# ==========================================
# NUEVA FUNCIÓN ETL: COSTOS MARGINALES (ZIP)
# ==========================================
@st.cache_data(show_spinner=False)
def extraer_cmg(fecha):
    año = fecha.strftime("%Y")
    mes_num = fecha.strftime("%m")
    dia = fecha.strftime("%d")
    mes_titulo = MESES[fecha.month]
    fecha_str = f"{año}{mes_num}{dia}"
    
    path_zip = f"Post Operación/Reportes/IEOD/{año}/{mes_num}_{mes_titulo}/{dia}/CMg{fecha_str}.zip"
    url_cmg = f"https://www.coes.org.pe/portal/browser/download?url={urllib.parse.quote(path_zip)}"
    
    headers = {'User-Agent': 'Mozilla/5.0'}
    df_cmg_dia = pd.DataFrame()
    
    try:
        res = requests.get(url_cmg, headers=headers, timeout=20)
        if res.status_code == 200:
            with zipfile.ZipFile(io.BytesIO(res.content)) as z:
                archivos_excel = [f for f in z.namelist() if f.endswith(('.xlsx', '.xls'))]
                if archivos_excel:
                    with z.open(archivos_excel[0]) as f:
                        xls = pd.ExcelFile(f, engine='openpyxl')
                        if 'Cmg_Barra' in xls.sheet_names:
                            df_raw = pd.read_excel(xls, sheet_name='Cmg_Barra', header=None)
                            
                            barras = df_raw.iloc[2, :].values
                            data = df_raw.iloc[3:51, :].values
                            
                            df_temp = pd.DataFrame(data, columns=barras)
                            fechas_horas = [fecha + timedelta(minutes=30 * (i + 1)) for i in range(48)]
                            df_temp['FECHA_HORA'] = fechas_horas
                            
                            barras_objetivo = ["SANTA ROSA 220", "MONTALVO 220", "TRUJILLO 220"]
                            cols_presentes = [c for c in barras_objetivo if c in df_temp.columns]
                            
                            if cols_presentes:
                                df_cmg_dia = df_temp[['FECHA_HORA'] + cols_presentes].copy()
                                df_cmg_dia = df_cmg_dia.melt(id_vars=['FECHA_HORA'], var_name='BARRA', value_name='CMG_USD')
                                df_cmg_dia['CMG_USD'] = pd.to_numeric(df_cmg_dia['CMG_USD'], errors='coerce').fillna(0)
    except Exception as e:
        pass
        
    return df_cmg_dia

def procesar_rango_fechas(start_date, end_date, progress_bar, status_text, dict_recursos):
    fechas = pd.date_range(start_date, end_date)
    total_dias = len(fechas)
    lista_dfs, lista_inter, lista_seguridad, lista_demanda, lista_cmg, alertas = [], [], [], [], [], []
    
    for i, f in enumerate(fechas):
        status_text.markdown(f"**⏳ Sincronizando datos de Despacho (COES):** {f.strftime('%d/%m/%Y')} *(Día {i+1} de {total_dias})*")
        df_dia, df_inter_dia, df_seguridad_dia, df_demanda_dia, error = extraer_datos_despacho(f, dict_recursos)
        df_cmg_dia = extraer_cmg(f)
        
        if not df_dia.empty: lista_dfs.append(df_dia)
        if not df_inter_dia.empty: lista_inter.append(df_inter_dia)
        if not df_seguridad_dia.empty: lista_seguridad.append(df_seguridad_dia)
        if not df_demanda_dia.empty: lista_demanda.append(df_demanda_dia)
        if not df_cmg_dia.empty: lista_cmg.append(df_cmg_dia)
        if error: alertas.append(error)
            
        progress_bar.progress((i + 1) / total_dias)
            
    df_final = pd.concat(lista_dfs, ignore_index=True) if lista_dfs else pd.DataFrame()
    df_inter_final = pd.concat(lista_inter, ignore_index=True) if lista_inter else pd.DataFrame()
    df_seguridad_final = pd.concat(lista_seguridad, ignore_index=True) if lista_seguridad else pd.DataFrame()
    df_demanda_final = pd.concat(lista_demanda, ignore_index=True) if lista_demanda else pd.DataFrame()
    df_cmg_final = pd.concat(lista_cmg, ignore_index=True) if lista_cmg else pd.DataFrame()
    
    return df_final, df_inter_final, df_seguridad_final, df_demanda_final, df_cmg_final, alertas


# --- 3. INTERFAZ DE USUARIO ---
st.sidebar.header("⚙️ Parámetros del Dashboard SEIN")

# Fechas por defecto: Día de hoy
hoy = datetime.now().date()
rango_fechas = st.sidebar.date_input("Intervalo de Fechas", value=(hoy, hoy))

if st.sidebar.button("📊 Extraer Despacho - SEIN Completo", type="primary"):
    if isinstance(rango_fechas, tuple) and len(rango_fechas) == 2:
        start_date, end_date = rango_fechas
        status_text = st.empty()
        progress_bar = st.progress(0)
        
        df_consolidado, df_inter_consolidado, df_seg_consolidado, df_dem_consolidado, df_cmg_consolidado, alertas = procesar_rango_fechas(start_date, end_date, progress_bar, status_text, dict_recursos_maestro)
        
        st.session_state['df_despacho'] = df_consolidado
        st.session_state['df_interconexiones'] = df_inter_consolidado
        st.session_state['df_seguridad'] = df_seg_consolidado
        st.session_state['df_demanda'] = df_dem_consolidado
        st.session_state['df_cmg'] = df_cmg_consolidado
        st.session_state['alertas_despacho'] = alertas
            
        status_text.empty()
        progress_bar.empty()


# --- 4. VISUALIZACIÓN DE DATOS ---
if 'df_despacho' in st.session_state:
    df_raw = st.session_state['df_despacho']
    df_inter_raw = st.session_state.get('df_interconexiones', pd.DataFrame())
    df_seg_raw = st.session_state.get('df_seguridad', pd.DataFrame())
    df_dem_raw = st.session_state.get('df_demanda', pd.DataFrame())
    alertas = st.session_state['alertas_despacho']
    
    if df_raw.empty:
        st.error("🚨 Extracción fallida o sin datos operacionales.")
    elif df_matriz_centrales.empty:
        st.error("❌ No se pudo cargar la matriz de centrales CetralesSEIN.xlsx")
    else:
        # ==========================================
        # FILTROS DINÁMICOS EN CASCADA
        # ==========================================
        st.sidebar.markdown("---")
        st.sidebar.header("🔍 Filtros Operativos")
        
        # Filtro 1: Área
        opts_zona = sorted([x for x in df_matriz_centrales['AREA_OPERATIVA'].unique() if x])
        filtro_zona = st.sidebar.multiselect("📍 Área Operativa:", options=opts_zona, placeholder="Todas")
        df_f1 = df_matriz_centrales[df_matriz_centrales['AREA_OPERATIVA'].isin(filtro_zona)] if filtro_zona else df_matriz_centrales

        # Filtro 2: Tipo Integrante (COES/NO COES)
        opts_int = sorted([x for x in df_f1['TIPO_INTEGRANTE'].unique() if x])
        defecto_int = ["COES"] if "COES" in opts_int else []
        filtro_int = st.sidebar.multiselect("⚖️ Tipo Integrante:", options=opts_int, default=defecto_int, placeholder="Todos")
        df_f2 = df_f1[df_f1['TIPO_INTEGRANTE'].isin(filtro_int)] if filtro_int else df_f1
        
        # Filtro 3: Req. Especial
        opts_req = sorted([x for x in df_f2['REQUERIMIENTO_ESPECIAL'].unique() if x])
        filtro_req = st.sidebar.multiselect("⚠️ Req. Especial:", options=opts_req, placeholder="Todos")
        df_f3 = df_f2[df_f2['REQUERIMIENTO_ESPECIAL'].isin(filtro_req)] if filtro_req else df_f2

        # Filtro 4: Empresa
        opts_emp = sorted([x for x in df_f3['EMPRESA_DESPACHO'].unique() if x])
        filtro_empresa = st.sidebar.multiselect("🏢 Empresa:", options=opts_emp, placeholder="Todas")
        df_f4 = df_f3[df_f3['EMPRESA_DESPACHO'].isin(filtro_empresa)] if filtro_empresa else df_f3

        # Filtro 5: Tipo de Generación (Basado en la Columna H de tu maestro)
        opts_tipo = sorted([x for x in df_f4['TIPO_GENERACION'].unique() if x])
        filtro_tipo = st.sidebar.multiselect("⚡ Tipo de Recurso:", options=opts_tipo, placeholder="Todas")
        df_f5 = df_f4[df_f4['TIPO_GENERACION'].isin(filtro_tipo)] if filtro_tipo else df_f4

        # Filtro 6: Central
        opts_cen = sorted([x for x in df_f5['CENTRAL'].unique() if x])
        filtro_cen = st.sidebar.multiselect("🏭 Central:", options=opts_cen, placeholder="Todas")
        df_f_final = df_f5[df_f5['CENTRAL'].isin(filtro_cen)] if filtro_cen else df_f5

        # Consolidar listas finales para enlace
        centrales_filtradas = df_f_final['CENTRAL'].str.strip().str.upper().tolist()
        nombres_calificacion_activos = [str(n).strip().upper() for n in df_f_final['CENTRAL_CALIFICACION'].unique() if n and n != "N/A" and str(n).lower() != 'nan']

        # Enlace ultra-robusto del Filtro con df_raw
        def es_central_valida(central_name):
            # Elimina sufijos como "(BIO)", "(HID)", "(GAS)" generados en ETL
            nom_limpio = re.sub(r'\s*\([^)]*\)$', '', str(central_name)).strip().upper()
            
            # Validación flexible de coincidencia parcial
            for c_filt in centrales_filtradas:
                if c_filt == nom_limpio or c_filt in nom_limpio or nom_limpio in c_filt:
                    return True
            return False

        df_datos = df_raw[df_raw['CENTRAL'].apply(es_central_valida)]
        
        if df_datos.empty:
            st.warning("⚠️ No hay datos despachados para las centrales filtradas en las fechas seleccionadas.")
        else:
            st.success("✅ Datos mostrados correctamente.")
            
            # Definición centralizada de colores para usar en todo el dashboard (ACTUALIZADO SEGÚN ESPECIFICACIÓN)
            colores_tecnologia = {
                "EÓLICA": "#808080", "EOLICA": "#808080",       # Plomo
                "HIDROELÉCTRICA": "#00BFFF", "HIDROELECTRICA": "#00BFFF", # Celeste
                "SOLAR": "#FFD700",                             # Amarillo
                "DIESEL/RESIDUAL": "#FF0000",                   # Rojo
                "GAS DE LA SELVA": "#90EE90",                   # Verde Claro
                "GAS CAMISEA": "#006400",                       # Verde Oscuro
                "GAS NORTE": "#0bb613",                         # Verde Más Oscuro
                "BIOMASA": "#800080",                           # Púrpura
                "GAS": "#006400"                                # Fallback genérico Verde Oscuro
            }
            
            # --- CONTENEDOR PARA EL BOTÓN DE REPORTE ---
            contenedor_reporte = st.container()
            
            es_formato_anexo1 = (df_datos['EMPRESA'] == 'N/A').all()
            st.markdown("### 📊 Datos Seleccionados (SEIN )")
            
            df_datos['FECHA_DIA'] = df_datos['FECHA_HORA'].dt.date

            if df_datos.empty:
                st.warning("⚠️ No hay datos despachados para la selección actual.")
            else:
                # INICIALIZACIÓN DE VARIABLES PARA EL REPORTE
                fig_tipo = None
                fig_inter = None
                fig_cen = None
                fig_prom = None
                fig_inactividad = None
                fig_actividad = None
                fig_bar_seg = None
                df_bar_data = pd.DataFrame()
                
                # FUNCIÓN HELPER PARA GRÁFICAS DE ÁREA (SIN DEMANDA)
                def crear_grafica_area(df_grafico, col_color, titulo, color_map=None):
                    df_plot = df_grafico.copy().dropna(subset=[col_color])
                    df_plot['DESPACHO_MW'] = pd.to_numeric(df_plot['DESPACHO_MW'], errors='coerce').fillna(0)
                    
                    df_sistema = df_plot.groupby('FECHA_HORA', as_index=False)['DESPACHO_MW'].sum()
                    max_demanda_real = df_sistema['DESPACHO_MW'].max()
                    
                    # Ampliamos el margen a 1.12 para que no se corte el texto del marcador Máximo
                    limite_superior_y = max_demanda_real * 1.12 if pd.notna(max_demanda_real) and max_demanda_real > 0 else 1000
                    fecha_min, fecha_max = df_plot['FECHA_HORA'].min(), df_plot['FECHA_HORA'].max()

                    fig = px.area(
                        df_plot, x="FECHA_HORA", y="DESPACHO_MW", color=col_color, 
                        title=titulo, labels={col_color: "Clasificación"}, 
                        color_discrete_map=color_map,
                        color_discrete_sequence=px.colors.qualitative.Alphabet,
                        template="plotly_white"
                    )
                    
                    fig.update_traces(hovertemplate="%{y:,.2f} MW", line=dict(width=0))
                    
                    # Línea de Total Generación (Transparente para el hover)
                    fig.add_scatter(
                        x=df_sistema['FECHA_HORA'], y=df_sistema['DESPACHO_MW'], mode='lines',
                        line=dict(width=0, color='rgba(0,0,0,0)'), name='<b>⚡ TOTAL GENERACIÓN</b>',
                        hovertemplate='<b>🗓️ %{x|%d/%m/%Y %H:%M} ➡️ %{y:,.2f} MW</b>', showlegend=False
                    )
                    
                    
                    # --- MARCADORES DE MÁXIMO Y MÍNIMO PARA LA CURVA TOTAL ---
                    if not df_sistema.empty:
                        idx_max = df_sistema['DESPACHO_MW'].idxmax()
                        idx_min = df_sistema['DESPACHO_MW'].idxmin()
                        
                        max_row = df_sistema.loc[idx_max]
                        min_row = df_sistema.loc[idx_min]
                        
                        # Marcador Máximo
                        fig.add_scatter(
                            x=[max_row['FECHA_HORA']], y=[max_row['DESPACHO_MW']],
                            mode='markers+text', marker=dict(color='black', size=12, symbol='triangle-up'),
                            text=[f"<b>Máx: {max_row['DESPACHO_MW']:,.0f} MW</b>"], textposition="top center",
                            textfont=dict(color="blue"),
                            name='Máx Total', hoverinfo='skip', showlegend=False
                        )
                        
                        # Marcador Mínimo
                        fig.add_scatter(
                            x=[min_row['FECHA_HORA']], y=[min_row['DESPACHO_MW']],
                            mode='markers+text', marker=dict(color='black', size=12, symbol='triangle-down'),
                            text=[f"<b>Mín: {min_row['DESPACHO_MW']:,.0f} MW</b>"], textposition="bottom center",
                            textfont=dict(color="blue"),
                            name='Mín Total', hoverinfo='skip', showlegend=False
                        )    
                    
                    fig.update_layout(
                        hovermode="x unified",
                        xaxis=dict(tickformat="%d/%m\n%H:%M", title="Fecha Operativa", range=[fecha_min, fecha_max]),
                        yaxis=dict(title="Potencia Activa (MW)", range=[0, limite_superior_y]),
                        height=550, margin=dict(t=50, b=50, l=50, r=20) 
                    )
                    return fig

                # PRECALCULO PARA GRÁFICAS DE UNIDADES
                energia_total_cen = df_datos.groupby('CENTRAL')['DESPACHO_MW'].sum()
                centrales_activas = energia_total_cen[energia_total_cen > 0].index
                
                df_plot_cen = df_datos[df_datos['CENTRAL'].isin(centrales_activas)].copy()
                energia_ordenada_cen = energia_total_cen[centrales_activas].sort_values(ascending=False).index
                df_plot_cen['CENTRAL'] = pd.Categorical(df_plot_cen['CENTRAL'], categories=energia_ordenada_cen, ordered=True)
                df_plot_cen = df_plot_cen.sort_values(['FECHA_HORA', 'CENTRAL'])

                # ==========================================
                # 1. DESPACHO POR TIPO DE GENERACIÓN (SEIN)
                # ==========================================
                st.markdown("---")
                st.header("1. 🏭 Despacho por Tipo de Generación (SEIN)")
                if not es_formato_anexo1:
                    df_tipo = df_datos.groupby(['FECHA_HORA', 'TIPO_CENTRAL'], as_index=False)['DESPACHO_MW'].sum()
                    energia_tipo = df_tipo.groupby('TIPO_CENTRAL')['DESPACHO_MW'].sum().sort_values(ascending=False).index
                    
                    orden_apilamiento = [
                        "BIOMASA", 
                        "SOLAR", 
                        "EÓLICA", "EOLICA", 
                        "HIDROELÉCTRICA", "HIDROELECTRICA", 
                        "GAS NORTE", 
                        "GAS DE LA SELVA", 
                        "GAS CAMISEA", 
                        "DIESEL/RESIDUAL"
                    ]

                    df_tipo['TIPO_CENTRAL'] = pd.Categorical(df_tipo['TIPO_CENTRAL'], categories=orden_apilamiento, ordered=True)
                    df_tipo = df_tipo.sort_values(['FECHA_HORA', 'TIPO_CENTRAL'])

                    fig_tipo = crear_grafica_area(df_tipo, 'TIPO_CENTRAL', "Curva Apilada por Tecnología - SEIN (MW)", color_map=colores_tecnologia)
                    st.plotly_chart(fig_tipo, use_container_width=True)
                else:
                    st.info("La vista por Tipo de Generación requiere el formato AnexoA (con metadatos), el archivo actual no lo contiene.")

                
                # ==========================================
                # 2. COSTOS MARGINALES (CMg) EN BARRAS DE REFERENCIA
                # ==========================================
                st.markdown("---")
                st.header("2. 💸 Despacho Operativo vs Costo Marginal (CMg)")
                st.info("Despacho por tecnología y flujos (Eje Izquierdo - MW) junto al Costo Marginal de la barra de referencia de Trujillo (Eje Derecho - S/./MWh). Curves de líneas resaltadas en negrita.")
                
                df_cmg_raw = st.session_state.get('df_cmg', pd.DataFrame())
                
                if df_cmg_raw.empty:
                    st.info("No se encontraron datos de Costos Marginales en el periodo descargado.")
                else:
                    import plotly.graph_objects as go
                    from plotly.subplots import make_subplots
                    
                    df_cmg_plot = df_cmg_raw.sort_values(['FECHA_HORA', 'BARRA']).copy()
                    
                    # FILTRO: Mantener solo la barra de TRUJILLO 220
                    df_cmg_plot = df_cmg_plot[df_cmg_plot['BARRA'] == 'TRUJILLO 220']
                    
                    # 1. Recuperar y procesar flujos negativos de la Sección 2
                    df_inter_raw = st.session_state.get('df_interconexiones', pd.DataFrame())
                    df_cn_total = pd.DataFrame()
                    df_l5006_total = pd.DataFrame()
                    
                    if not df_inter_raw.empty:
                        # Negativo del Flujo Neto Centro-Norte
                        df_cn = df_inter_raw[df_inter_raw['ENLACE'] == 'CENTRO-NORTE'].copy()
                        if not df_cn.empty:
                            df_cn_total = df_cn.groupby('FECHA_HORA', as_index=False)['FLUJO_MW'].sum()
                            df_cn_total['FLUJO_NEG'] = df_cn_total['FLUJO_MW'] * -1
                        
                        # Negativo del Flujo de la línea específica L-5006
                        df_l5006 = df_inter_raw[df_inter_raw['LINEA_TRANSMISION'].str.contains('L-5006', case=False, na=False)].copy()
                        if not df_l5006.empty:
                            df_l5006_total = df_l5006.groupby('FECHA_HORA', as_index=False)['FLUJO_MW'].sum()
                            df_l5006_total['FLUJO_NEG'] = df_l5006_total['FLUJO_MW'] * -1
                    
                    # Paleta de colores para la barra de Trujillo (Azul claro)
                    colores_barra = {
                        'TRUJILLO 220': "#0099FF"  # Azul claro
                    }
                    
                    fig_cmg = make_subplots(specs=[[{"secondary_y": True}]])
                    
                    # --- A. DESPACHO POR TECNOLOGÍA (EJE Y PRIMARIO - MW) ---
                    df_tipo_cmg = df_datos.groupby(['FECHA_HORA', 'TIPO_CENTRAL'], as_index=False)['DESPACHO_MW'].sum()
                    orden_apilamiento = [
                        "BIOMASA", "SOLAR", "EÓLICA", "EOLICA", 
                        "HIDROELÉCTRICA", "HIDROELECTRICA", 
                        "GAS NORTE", "GAS DE LA SELVA", "GAS CAMISEA", 
                        "DIESEL/RESIDUAL"
                    ]
                    
                    for tec in orden_apilamiento:
                        df_tec = df_tipo_cmg[df_tipo_cmg['TIPO_CENTRAL'] == tec]
                        if not df_tec.empty:
                            fig_cmg.add_trace(
                                go.Scatter(
                                    x=df_tec['FECHA_HORA'], y=df_tec['DESPACHO_MW'],
                                    mode='lines', line=dict(width=0),
                                    fill='tonexty', stackgroup='one',
                                    name=tec, marker_color=colores_tecnologia.get(tec, '#808080'),
                                    hovertemplate=f"<b>{tec}</b>: %{{y:,.2f}} MW"
                                ),
                                secondary_y=False
                            )
                            
                    # --- B. ENLACES DE FLUJO NEGATIVO EN NEGRITA Y PUNTEADOS (EJE Y PRIMARIO - MW) ---
                    if not df_cn_total.empty:
                        fig_cmg.add_trace(
                            go.Scatter(
                                x=df_cn_total['FECHA_HORA'], y=df_cn_total['FLUJO_NEG'],
                                mode='lines', line=dict(width=3, dash='dash', color='#9467bd'), # Negrita (width=3) y Punteada
                                name='⚡ FLUJO CENTRO-NORTE', marker_color='#9467bd',
                                hovertemplate="<b>FLUJO CENTRO-NORTE</b>: %{y:,.2f} MW"
                            ),
                            secondary_y=False
                        )
                        # Marcadores Máx/Mín Flujo Centro-Norte
                        idx_max_cn = df_cn_total['FLUJO_NEG'].idxmax()
                        idx_min_cn = df_cn_total['FLUJO_NEG'].idxmin()
                        max_cn = df_cn_total.loc[idx_max_cn]
                        min_cn = df_cn_total.loc[idx_min_cn]
                        
                        fig_cmg.add_trace(go.Scatter(
                            x=[max_cn['FECHA_HORA']], y=[max_cn['FLUJO_NEG']],
                            mode='markers+text', marker=dict(color='#9467bd', size=12, symbol='triangle-up'),
                            text=[f"<b>Máx C-N: {max_cn['FLUJO_NEG']:,.0f} MW</b>"], textposition="top center",
                            textfont=dict(color="blue"), # <--- Etiqueta en azul
                            showlegend=False, hoverinfo='skip'
                        ), secondary_y=False)
                        
                    if not df_l5006_total.empty:
                        fig_cmg.add_trace(
                            go.Scatter(
                                x=df_l5006_total['FECHA_HORA'], y=df_l5006_total['FLUJO_NEG'],
                                mode='lines', line=dict(width=3, dash='dot', color='#e377c2'), # Negrita (width=3) y Punteada
                                name='⚡ FLUJO L-5006',
                                hovertemplate="<b>FLUJO L-5006</b>: %{y:,.2f} MW"
                            ),
                            secondary_y=False
                        )

                        # Marcadores Máx/Mín Línea L-5006
                        idx_max_l = df_l5006_total['FLUJO_NEG'].idxmax()
                        idx_min_l = df_l5006_total['FLUJO_NEG'].idxmin()
                        max_l = df_l5006_total.loc[idx_max_l]
                        min_l = df_l5006_total.loc[idx_min_l]
                        
                        fig_cmg.add_trace(go.Scatter(
                            x=[max_l['FECHA_HORA']], y=[max_l['FLUJO_NEG']],
                            mode='markers+text', marker=dict(color='#e377c2', size=12, symbol='triangle-up'),
                            text=[f"<b>Máx L-5006: {max_l['FLUJO_NEG']:,.0f} MW</b>"], textposition="top center",
                            textfont=dict(color="blue"), # <--- Etiqueta en azul
                            showlegend=False, hoverinfo='skip'
                        ), secondary_y=False)

                    # --- C. LÍNEAS PUNTEADAS DE COSTO MARGINAL EN NEGRITA (EJE Y SECUNDARIO - S/./MWh) ---
                    def graficar_min_max_cmg_dual(fig, df_filtro, color_marcador, nombre_barra):
                        if not df_filtro.empty:
                            idx_max = df_filtro['CMG_USD'].idxmax()
                            idx_min = df_filtro['CMG_USD'].idxmin()
                            
                            fig.add_trace(
                                go.Scatter(
                                    x=[df_filtro.loc[idx_max, 'FECHA_HORA']], y=[df_filtro.loc[idx_max, 'CMG_USD']],
                                    mode='markers+text', marker=dict(color=color_marcador, size=12, symbol='triangle-up'),
                                    text=[f"<b>Máx: {df_filtro.loc[idx_max, 'CMG_USD']:,.1f} S/./MWh</b>"], textposition="top center",
                                    textfont=dict(color="blue"), # <--- Etiqueta en azul
                                    name=f'Máx {nombre_barra}', hoverinfo='skip', showlegend=False
                                ),
                                secondary_y=True
                            )

                    for barra in df_cmg_plot['BARRA'].unique():
                        df_barra = df_cmg_plot[df_cmg_plot['BARRA'] == barra]
                        fig_cmg.add_trace(
                            go.Scatter(
                                x=df_barra['FECHA_HORA'], y=df_barra['CMG_USD'],
                                mode='lines', line=dict(width=3, dash='dot', color=colores_barra.get(barra, '#66b3ff')), # Azul claro y negrita
                                name=barra, hovertemplate=f"<b>{barra}</b>: %{{y:,.2f}} S/./MWh"
                            ),
                            secondary_y=True
                        )
                        graficar_min_max_cmg_dual(fig_cmg, df_barra, colores_barra.get(barra, '#66b3ff'), barra)
                    
                    # --- D. NUEVO: LÍMITES DE TRANSMISIÓN DE LA LÍNEA L-5006 (EJE Y PRIMARIO) ---
                    fecha_min_cmg = df_cmg_plot['FECHA_HORA'].min()
                    fecha_max_cmg = df_cmg_plot['FECHA_HORA'].max()
                    
                    # Límite Superior: 700 MW (Línea punteada con marcador al final / lado derecho)
                    fig_cmg.add_shape(type="line", x0=fecha_min_cmg, y0=700, x1=fecha_max_cmg, y1=700,
                                      line=dict(color="red", width=2, dash="dash"), yref="y")
                    fig_cmg.add_annotation(x=fecha_max_cmg, y=700, text="<b>Límite Sup L-5006: 700 MW</b>",
                                           showarrow=False, xanchor="right", yanchor="bottom", yref="y",
                                           font=dict(color="blue")) # <--- Etiqueta en azul
                                           
                    # Límite Inferior: 600 MW (Línea punteada con marcador al final / lado derecho)
                    fig_cmg.add_shape(type="line", x0=fecha_min_cmg, y0=600, x1=fecha_max_cmg, y1=600,
                                      line=dict(color="green", width=2, dash="dash"), yref="y")
                    fig_cmg.add_annotation(x=fecha_max_cmg, y=600, text="<b>Límite Inf L-5006: 600 MW</b>",
                                           showarrow=False, xanchor="right", yanchor="bottom", yref="y",
                                           font=dict(color="blue")) # <--- Etiqueta en azul

                    # --- E. AJUSTE DE RANGOS DE EJES Y DISEÑO ---
                    # Cálculo del límite superior agregando exactamente 400 MW de holgura visual
                    if not df_tipo_cmg.empty:
                        max_gen_total = df_tipo_cmg.groupby('FECHA_HORA')['DESPACHO_MW'].sum().max()
                        max_flujo_neg = max(df_cn_total['FLUJO_NEG'].max() if not df_cn_total.empty else 0,
                                            df_l5006_total['FLUJO_NEG'].max() if not df_l5006_total.empty else 0)
                        limite_y1 = max(max_gen_total, max_flujo_neg) + 400
                    else:
                        limite_y1 = 1400
                    
                    max_val_cmg = df_cmg_plot['CMG_USD'].max()
                    limite_y2 = max_val_cmg * 1.15 if max_val_cmg > 0 else 50
                    
                    fig_cmg.update_layout(
                        hovermode="x unified",
                        height=650, 
                        margin=dict(t=50, b=50, l=50, r=150), # Se amplía el margen derecho para la leyenda
                        legend=dict(
                            title="<b>Componentes SEIN</b>",
                            orientation="v",         # Leyenda en forma vertical
                            yanchor="top", 
                            y=1, 
                            xanchor="left", 
                            x=1.05                   # Posicionada a la derecha de los ejes
                        ),
                        template="plotly_white"
                    )
                    
                    fig_cmg.update_xaxes(title_text="Fecha Operativa", tickformat="%d/%m\n%H:%M")
                    fig_cmg.update_yaxes(title_text="Potencia Activa (MW)", range=[0, limite_y1], secondary_y=False)
                    fig_cmg.update_yaxes(title_text="Costo Marginal (S/./MWh)", range=[0, limite_y2], secondary_y=True, showgrid=False)
                    
                    st.plotly_chart(fig_cmg, use_container_width=True)
                    
                    with st.expander("Ver Datos de CMg (Vista Matricial)"):
                        df_cmg_pivot = df_cmg_plot.copy()
                        df_cmg_pivot['FECHA'] = df_cmg_pivot['FECHA_HORA'].dt.strftime('%d/%m/%Y')
                        df_cmg_pivot['HORA'] = df_cmg_pivot['FECHA_HORA'].dt.strftime('%H:%M')
                        matriz_cmg = df_cmg_pivot.pivot_table(index=['FECHA', 'HORA'], columns=['BARRA'], values='CMG_USD', aggfunc='mean').round(2)
                        st.dataframe(matriz_cmg, use_container_width=True)

                # ==========================================
                # 3. FLUJO DE ENLACES (CENTRO-NORTE Y CENTRO-SUR)
                # ==========================================
                st.markdown("---")
                st.header("3. 🔌 Flujo de Enlaces")
                
                if df_inter_raw.empty:
                    st.info("No se encontraron datos de enlaces en el periodo descargado.")
                else:
                    df_inter_plot = df_inter_raw.sort_values(['FECHA_HORA', 'LINEA_TRANSMISION']).copy()
                    
                    # Función auxiliar optimizada con etiquetas de texto en color azul
                    def marcar_min_max_flujo(fig, df_total, color_marcador):
                        if not df_total.empty:
                            # 1. Encontrar los índices donde se registran la mayor y menor magnitud absoluta
                            idx_max_abs = df_total['FLUJO_MW'].abs().idxmax()
                            idx_min_abs = df_total['FLUJO_MW'].abs().idxmin()
                            
                            # 2. Extraer los valores con su signo original del COES
                            max_val = df_total.loc[idx_max_abs, 'FLUJO_MW']
                            min_val = df_total.loc[idx_min_abs, 'FLUJO_MW']
                            
                            # Recta Horizontal de Máxima Magnitud Absoluta
                            fig.add_hline(
                                y=max_val, 
                                line_dash="dash", 
                                line_color=color_marcador, 
                                line_width=2,
                                annotation_text=f"<b>Máx: {max_val:,.0f} MW</b>", 
                                annotation_position="top left",
                                annotation_font=dict(color="blue") # <--- MODIFICACIÓN: Color azul para la etiqueta
                            )
                            
                            # Recta Horizontal de Mínima Magnitud Absoluta
                            fig.add_hline(
                                y=min_val, 
                                line_dash="dash", 
                                line_color=color_marcador, 
                                line_width=2,
                                annotation_text=f"<b>Mín: {min_val:,.0f} MW</b>", 
                                annotation_position="bottom left",
                                annotation_font=dict(color="blue") # <--- MODIFICACIÓN: Color azul para la etiqueta
                            )
                            
                            # 3. Autoajuste dinámico del eje Y delegado al motor gráfico nativo
                            fig.update_layout(yaxis=dict(autorange=True))
                    
                    # GRAFICAR CENTRO-NORTE (GRÁFICA SUPERIOR)
                    df_cn = df_inter_plot[df_inter_plot['ENLACE'] == 'CENTRO-NORTE'].copy()
                    if not df_cn.empty:
                        fig_inter_cn = px.area(
                            df_cn, x="FECHA_HORA", y="FLUJO_MW", color="LINEA_TRANSMISION",
                            title="Flujo de Potencia - Enlace Centro-Norte (MW)",
                            labels={"LINEA_TRANSMISION": "Línea", "FLUJO_MW": "Potencia (MW)"},
                            color_discrete_sequence=['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd'],
                            template="plotly_white"
                        )
                        df_cn_total = df_cn.groupby('FECHA_HORA', as_index=False)['FLUJO_MW'].sum()
                        fig_inter_cn.update_traces(hovertemplate="%{y:,.2f} MW", line=dict(width=0))
                        fig_inter_cn.add_scatter(
                            x=df_cn_total['FECHA_HORA'], y=df_cn_total['FLUJO_MW'], mode='lines',
                            line=dict(width=3, color='gray', dash='solid'), name='<b>⚡ TOTAL C-N</b>',
                            hovertemplate='<b>🗓️ %{x|%d/%m/%Y %H:%M} ➡️ %{y:,.2f} MW</b>'
                        )
                        
                        # Invocación de la función modificada
                        marcar_min_max_flujo(fig_inter_cn, df_cn_total, 'black')
                        
                        fig_inter_cn.update_layout(hovermode="x unified", xaxis_title="Fecha Operativa", height=450)
                        st.plotly_chart(fig_inter_cn, use_container_width=True)

                    # GRAFICAR CENTRO-SUR (GRÁFICA INFERIOR)
                    df_cs = df_inter_plot[df_inter_plot['ENLACE'] == 'CENTRO-SUR'].copy()
                    if not df_cs.empty:
                        fig_inter_cs = px.area(
                            df_cs, x="FECHA_HORA", y="FLUJO_MW", color="LINEA_TRANSMISION",
                            title="Flujo de Potencia - Enlace Centro-Sur (MW)",
                            labels={"LINEA_TRANSMISION": "Línea", "FLUJO_MW": "Potencia (MW)"},
                            color_discrete_sequence=['#8c564b', '#e377c2', '#7f7f7f', '#bcbd22'],
                            template="plotly_white"
                        )
                        df_cs_total = df_cs.groupby('FECHA_HORA', as_index=False)['FLUJO_MW'].sum()
                        fig_inter_cs.update_traces(hovertemplate="%{y:,.2f} MW", line=dict(width=0))
                        fig_inter_cs.add_scatter(
                            x=df_cs_total['FECHA_HORA'], y=df_cs_total['FLUJO_MW'], mode='lines',
                            line=dict(width=3, color='gray', dash='solid'), name='<b>⚡ TOTAL C-S</b>',
                            hovertemplate='<b>🗓️ %{x|%d/%m/%Y %H:%M} ➡️ %{y:,.2f} MW</b>'
                        )
                        
                        # Invocación de la función modificada
                        marcar_min_max_flujo(fig_inter_cs, df_cs_total, 'black')
                        
                        fig_inter_cs.update_layout(hovermode="x unified", xaxis_title="Fecha Operativa", height=450)
                        st.plotly_chart(fig_inter_cs, use_container_width=True)
                    
                    with st.expander("Ver Datos de Enlaces (Vista Matricial)"):
                        df_inter_pivot = df_inter_plot.copy()
                        df_inter_pivot['FECHA'] = df_inter_pivot['FECHA_HORA'].dt.strftime('%d/%m/%Y')
                        df_inter_pivot['HORA'] = df_inter_pivot['FECHA_HORA'].dt.strftime('%H:%M')
                        
                        df_mat_inter = df_inter_pivot.pivot_table(
                            index=['FECHA', 'HORA'], columns=['ENLACE', 'LINEA_TRANSMISION'], values='FLUJO_MW', aggfunc='sum'
                        ).round(2).fillna(0)
                        st.dataframe(df_mat_inter, use_container_width=True)

                # ==========================================
                # 4. GENERACIÓN SEIN (SIN DEMANDA)
                # ==========================================
                st.markdown("---")
                st.header("4. 📊 Generación del SEIN por Central")
                
                df_plot_cen_aux = df_plot_cen.copy().dropna(subset=['CENTRAL'])
                df_plot_cen_aux['DESPACHO_MW'] = pd.to_numeric(df_plot_cen_aux['DESPACHO_MW'], errors='coerce').fillna(0)
                
                df_sistema_aux = df_plot_cen_aux.groupby('FECHA_HORA', as_index=False)['DESPACHO_MW'].sum()
                max_demanda_real_aux = df_sistema_aux['DESPACHO_MW'].max()
                
                limite_superior_y_aux = max_demanda_real_aux * 1.05 if pd.notna(max_demanda_real_aux) and max_demanda_real_aux > 0 else 1000
                fecha_min_aux, fecha_max_aux = df_plot_cen_aux['FECHA_HORA'].min(), df_plot_cen_aux['FECHA_HORA'].max()

                fig_cen = px.area(
                    df_plot_cen_aux, x="FECHA_HORA", y="DESPACHO_MW", color='CENTRAL', 
                    title="Despacho de Potencia por Unidad - SEIN (MW)", labels={'CENTRAL': "Central"}, 
                    color_discrete_sequence=['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', '#8c564b', '#e377c2', '#7f7f7f', '#bcbd22', '#17becf'],
                    template="plotly_white"
                )
                
                fig_cen.update_traces(hovertemplate="%{y:,.2f} MW", line=dict(width=0))
                fig_cen.add_scatter(
                    x=df_sistema_aux['FECHA_HORA'], y=df_sistema_aux['DESPACHO_MW'], mode='lines',
                    line=dict(width=0, color='rgba(0,0,0,0)'), name='<b>⚡ TOTAL GENERACIÓN</b>',
                    hovertemplate='<b>🗓️ %{x|%d/%m/%Y %H:%M} ➡️ %{y:,.2f} MW</b>', showlegend=False
                )
                
                fig_cen.update_layout(
                    hovermode="x unified",
                    xaxis=dict(tickformat="%d/%m\n%H:%M", title="Fecha Operativa", range=[fecha_min_aux, fecha_max_aux]),
                    yaxis=dict(title="Potencia Activa (MW)", range=[0, limite_superior_y_aux]),
                    height=550, margin=dict(t=50, b=50, l=50, r=20) 
                )
                st.plotly_chart(fig_cen, use_container_width=True)

                # ==========================================
                # 5. POTENCIA PROMEDIO DIARIA
                # ==========================================
                st.markdown("---")
                st.header("5. 📈 Potencia Promedio Diaria (SEIN)")
                
                # CORRECCIÓN DE EFECTO DE BORDE: Asignar las 00:00 al día operativo correcto restando 1 minuto
                df_plot_cen['FECHA_DIA_OPERATIVO'] = (df_plot_cen['FECHA_HORA'] - pd.Timedelta(minutes=1)).dt.date
                
                # Gráfica 5.1: Promedio de todo el día (24 Horas / 48 Periodos)
                df_promedio = df_plot_cen.groupby(['FECHA_DIA_OPERATIVO', 'CENTRAL'], as_index=False)['DESPACHO_MW'].mean()
                df_promedio['FECHA_DIA_OPERATIVO'] = pd.to_datetime(df_promedio['FECHA_DIA_OPERATIVO']).dt.strftime('%d/%m/%Y')
                
                fig_prom = px.bar(
                    df_promedio, x='FECHA_DIA_OPERATIVO', y='DESPACHO_MW', color='CENTRAL',
                    title="Potencia Promedio Diaria Total (24 Horas) (MW)",
                    barmode='group',
                    labels={'FECHA_DIA_OPERATIVO': 'Día Operativo', 'DESPACHO_MW': 'Potencia Promedio (MW)'},
                    color_discrete_sequence=['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', '#8c564b', '#e377c2', '#7f7f7f', '#bcbd22', '#17becf'],
                    template="plotly_white"
                )
                fig_prom.update_layout(xaxis=dict(type='category'), height=500)
                st.plotly_chart(fig_prom, use_container_width=True)

                # Gráfica 5.2: Promedio Operativo (Solo en periodos con inyección > 0 MW)
                df_solo_inyeccion = df_plot_cen[df_plot_cen['DESPACHO_MW'] > 0].copy()
                
                if not df_solo_inyeccion.empty:
                    df_promedio_iny = df_solo_inyeccion.groupby(['FECHA_DIA_OPERATIVO', 'CENTRAL'], as_index=False)['DESPACHO_MW'].mean()
                    df_promedio_iny['FECHA_DIA_OPERATIVO'] = pd.to_datetime(df_promedio_iny['FECHA_DIA_OPERATIVO']).dt.strftime('%d/%m/%Y')
                    
                    fig_prom_iny = px.bar(
                        df_promedio_iny, x='FECHA_DIA_OPERATIVO', y='DESPACHO_MW', color='CENTRAL',
                        title="Potencia Promedio en Operación (MW)",
                        barmode='group',
                        labels={'FECHA_DIA_OPERATIVO': 'Día Operativo', 'DESPACHO_MW': 'Promedio en Operación (MW)'},
                        color_discrete_sequence=['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', '#8c564b', '#e377c2', '#7f7f7f', '#bcbd22', '#17becf'],
                        template="plotly_white"
                    )
                    fig_prom_iny.update_layout(xaxis=dict(type='category'), height=500)
                    st.plotly_chart(fig_prom_iny, use_container_width=True)
                else:
                    fig_prom_iny = None
                    st.info("No se registraron periodos con inyección de potencia mayor a 0 MW para calcular el promedio operativo.")

                # ==========================================
                # 6. CONTROL DE TIEMPOS: INACTIVIDAD Y OPERACIÓN
                # ==========================================
                st.markdown("---")
                st.header("6. ⏱️ Control de Tiempos: Inactividad y Operación (SEIN)")
                
                # --- PREPARACIÓN DE BLOQUES GANTT ---
                df_gantt = df_datos.copy()
                df_gantt = df_gantt.sort_values(['CENTRAL', 'FECHA_HORA'])
                
                # Clasificar estado operativo
                df_gantt['ESTADO'] = np.where(df_gantt['DESPACHO_MW'] > 0, 'OPERANDO', 'INACTIVO')
                
                # Detectar cambios de estado lógico para crear bloques continuos de tiempo
                df_gantt['CAMBIO_ESTADO'] = (df_gantt['ESTADO'] != df_gantt['ESTADO'].shift(1)) | (df_gantt['CENTRAL'] != df_gantt['CENTRAL'].shift(1))
                df_gantt['BLOQUE'] = df_gantt['CAMBIO_ESTADO'].cumsum()
                
                # Agrupar para obtener el inicio y fin de cada bloque
                df_bloques = df_gantt.groupby(['CENTRAL', 'TIPO_CENTRAL', 'ESTADO', 'BLOQUE'], as_index=False).agg(
                    INICIO=('FECHA_HORA', 'min'),
                    FIN=('FECHA_HORA', 'max')
                )
                # Sumar 30 minutos al final para cerrar el intervalo de manera exacta
                df_bloques['FIN'] = df_bloques['FIN'] + pd.Timedelta(minutes=30)
                
                # Fijar límites del Eje X para todos los Gantt (Intervalo de evaluación completo)
                fecha_inicio_gantt = df_datos['FECHA_HORA'].min()
                fecha_fin_gantt = df_datos['FECHA_HORA'].max() + pd.Timedelta(minutes=30)
                
                # --- GANTT 1: INACTIVIDAD (EXCLUYENDO DIÉSEL) ---
                st.markdown("#### 🚥 Cronograma de Inactividad (Tecnologías Base y Renovables)")
                st.info("Visualización de los bloques horarios donde las unidades NO inyectaron potencia (0 MW), excluyendo intencionalmente a las térmicas Diésel/Residual.")
                
                # FILTRO CLAVE: Bloques inactivos excluyendo DIESEL/RESIDUAL
                df_bloques_inactivos = df_bloques[(df_bloques['ESTADO'] == 'INACTIVO') & (df_bloques['TIPO_CENTRAL'] != 'DIESEL/RESIDUAL')].copy()
                
                if df_bloques_inactivos.empty:
                    st.success("✅ No se detectaron periodos de inactividad para las unidades base/renovables seleccionadas.")
                else:
                    fig_gantt_inact = px.timeline(
                        df_bloques_inactivos, 
                        x_start="INICIO", 
                        x_end="FIN", 
                        y="CENTRAL", 
                        color="TIPO_CENTRAL",
                        hover_data={"INICIO": "|%d/%m/%Y %H:%M", "FIN": "|%d/%m/%Y %H:%M"},
                        color_discrete_map=colores_tecnologia,
                        template="plotly_white"
                    )
                    
                    fig_gantt_inact.update_yaxes(autorange="reversed")
                    fig_gantt_inact.update_layout(
                        xaxis=dict(tickformat="%d/%m\n%H:%M", title="Línea de Tiempo (Periodos Inactivos)", range=[fecha_inicio_gantt, fecha_fin_gantt]),
                        yaxis=dict(title="Unidad Generadora", dtick=1),
                        height=max(400, len(df_bloques_inactivos['CENTRAL'].unique()) * 22),
                        margin=dict(t=30, b=50, l=50, r=20),
                        legend=dict(title="Tecnología", orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1)
                    )
                    st.plotly_chart(fig_gantt_inact, use_container_width=True)

                # --- RESUMEN ACUMULADO (BARRAS INACTIVIDAD) ---
                st.markdown("#### 📊 Resumen Acumulado de Horas por Unidad")
                
                df_tiempos = df_datos.copy()
                df_tiempos['INACTIVO_HR'] = (df_tiempos['DESPACHO_MW'] == 0) * 0.5
                df_tiempos['ACTIVO_HR'] = (df_tiempos['DESPACHO_MW'] > 0) * 0.5
                df_resumen_tiempos = df_tiempos.groupby(['CENTRAL', 'TIPO_CENTRAL'], as_index=False)[['INACTIVO_HR', 'ACTIVO_HR']].sum()
                
                tipos_inactividad = ['HIDROELÉCTRICA', 'HIDROELECTRICA', 'EÓLICA', 'EOLICA', 'SOLAR', 'BIOMASA', 'GAS DE LA SELVA', 'GAS CAMISEA', 'GAS NORTE']
                df_inactividad = df_resumen_tiempos[df_resumen_tiempos['TIPO_CENTRAL'].isin(tipos_inactividad)].copy()
                
                if not df_inactividad.empty:
                    fig_inactividad = px.bar(
                        df_inactividad, x='CENTRAL', y='INACTIVO_HR', color='TIPO_CENTRAL',
                        title="Horas No Despachadas (Inactividad) por Central (Excluye Diésel)",
                        labels={'CENTRAL': 'Central Generadora', 'INACTIVO_HR': 'Horas Inactivas (h)', 'TIPO_CENTRAL': 'Tipo'},
                        color_discrete_map=colores_tecnologia,
                        template="plotly_white"
                    )
                    fig_inactividad.update_layout(xaxis={'categoryorder':'total descending'}, height=450)
                    st.plotly_chart(fig_inactividad, use_container_width=True)
                else:
                    st.info("No hay datos de las tecnologías seleccionadas para inactividad.")

                # --- GANTT 2: ACTIVIDAD DIESEL/RESIDUAL ---
                st.markdown("---")
                st.markdown("#### 🚨 Cronograma de Operación (Diésel / Residual)")
                st.info("Visualización de los bloques donde las centrales Diesel/Residual inyectaron energía a la red.")
                
                # FILTRO CLAVE: Bloques activos solo para DIESEL/RESIDUAL
                df_bloques_diesel = df_bloques[(df_bloques['ESTADO'] == 'OPERANDO') & (df_bloques['TIPO_CENTRAL'] == 'DIESEL/RESIDUAL')].copy()
                
                if df_bloques_diesel.empty:
                    st.success("✅ Las unidades Diésel/Residual seleccionadas no registraron inyección de energía en el periodo evaluado.")
                else:
                    fig_gantt_diesel = px.timeline(
                        df_bloques_diesel, 
                        x_start="INICIO", 
                        x_end="FIN", 
                        y="CENTRAL", 
                        color="TIPO_CENTRAL",
                        hover_data={"INICIO": "|%d/%m/%Y %H:%M", "FIN": "|%d/%m/%Y %H:%M"},
                        color_discrete_map=colores_tecnologia,
                        template="plotly_white"
                    )
                    
                    fig_gantt_diesel.update_yaxes(autorange="reversed")
                    fig_gantt_diesel.update_layout(
                        xaxis=dict(tickformat="%d/%m\n%H:%M", title="Línea de Tiempo (Periodos de Inyección)", range=[fecha_inicio_gantt, fecha_fin_gantt]),
                        yaxis=dict(title="Unidad Diésel/Residual", dtick=1),
                        height=max(300, len(df_bloques_diesel['CENTRAL'].unique()) * 22),
                        margin=dict(t=30, b=50, l=50, r=20),
                        showlegend=False # Oculto pues sabemos que es de tecnología Diésel
                    )
                    st.plotly_chart(fig_gantt_diesel, use_container_width=True)

                # --- RESUMEN ACUMULADO (BARRAS ACTIVIDAD DIESEL) ---
                tipos_actividad = ['DIESEL/RESIDUAL']
                df_actividad = df_resumen_tiempos[df_resumen_tiempos['TIPO_CENTRAL'].isin(tipos_actividad)].copy()
                
                if not df_actividad.empty:
                    # Filtramos para mostrar solo las barras de las unidades que efectivamente operaron > 0 hrs
                    df_actividad_plot = df_actividad[df_actividad['ACTIVO_HR'] > 0]
                    if not df_actividad_plot.empty:
                        fig_actividad = px.bar(
                            df_actividad_plot, x='CENTRAL', y='ACTIVO_HR', color='TIPO_CENTRAL',
                            title="Horas Totales de Operación Activa (Centrales Diésel/Residual)",
                            labels={'CENTRAL': 'Central Generadora', 'ACTIVO_HR': 'Horas de Operación (h)', 'TIPO_CENTRAL': 'Tipo'},
                            color_discrete_map=colores_tecnologia,
                            template="plotly_white"
                        )
                        fig_actividad.update_layout(xaxis={'categoryorder':'total descending'}, height=400)
                        st.plotly_chart(fig_actividad, use_container_width=True)
                    else:
                        st.info("No se registraron horas acumuladas de operación Diésel/Residual en la selección actual.")
                else:
                    st.info("No hay centrales Diésel/Residual presentes en la selección actual.")

                # ==========================================
                # 7. CALIFICACIÓN DE LA OPERACIÓN (ENLAZADA AL FILTRO)
                # ==========================================
                st.markdown("---")
                st.header("7. 🛡️ Calificación de la Operación")
                
                if df_seg_raw.empty:
                    st.info("No se registraron calificaciones de operación en la hoja CALIFICA_OPE_UG para el periodo consultado.")
                elif not nombres_calificacion_activos:
                    st.warning("Las centrales seleccionadas no poseen mapeo de Calificación de Operación en la matriz.")
                else:
                    df_bar_data = df_seg_raw.dropna(subset=['INICIO', 'FIN']).copy()
                    if not df_bar_data.empty:
                        # ENLACE: Utilizar únicamente las centrales sobrevivientes del filtro en cascada
                        valid_centrales_calif = nombres_calificacion_activos
                        
                        def es_central_valida_calif(nombre):
                            nom_limpio = re.sub(r'\s+', ' ', str(nombre).strip().upper())
                            for valid_c in valid_centrales_calif:
                                if valid_c in nom_limpio:
                                    return True
                            return False
                            
                        df_bar_data = df_bar_data[df_bar_data['CENTRAL'].apply(es_central_valida_calif)].copy()
                        
                        if not df_bar_data.empty:
                            df_bar_data['CENTRAL_GRUPO'] = df_bar_data['CENTRAL'].astype(str) + " - " + df_bar_data['GRUPO'].astype(str)
                            df_bar_data['HORAS_OPERACION'] = (df_bar_data['FIN'] - df_bar_data['INICIO']).dt.total_seconds() / 3600.0
                            df_bar_data['HORAS_OPERACION'] = df_bar_data['HORAS_OPERACION'].clip(lower=0)

                            df_agrupado = df_bar_data.groupby(['CENTRAL_GRUPO', 'TIPO_OPERACION'], as_index=False)['HORAS_OPERACION'].sum()

                            colores_operacion = {
                                "POR SEGURIDAD": "#8B0000",
                                "POR POTENCIA O ENERGIA": "#00BFFF",
                                "A MINIMA CARGA": "#32CD32",
                                "POR COGENERACION": "#FF8C00",
                                "POR RSF": "#FFD700",
                                "POR PRUEBAS": "#808080"
                            }

                            fig_bar_seg = px.bar(
                                df_agrupado,
                                x="HORAS_OPERACION",
                                y="CENTRAL_GRUPO",
                                color="TIPO_OPERACION",
                                orientation='h',
                                title="Horas Totales de Operación por Unidad y Tipo de Operación",
                                labels={
                                    "HORAS_OPERACION": "Horas Totales (h)",
                                    "CENTRAL_GRUPO": "Unidad Generadora",
                                    "TIPO_OPERACION": "Tipo de Operación"
                                },
                                color_discrete_map=colores_operacion,
                                template="plotly_white"
                            )

                            fig_bar_seg.update_layout(
                                yaxis=dict(title="Unidad Generadora", categoryorder="total ascending"),
                                xaxis=dict(title="Horas de Operación Totales"),
                                height=max(400, len(df_agrupado['CENTRAL_GRUPO'].unique()) * 40)
                            )
                            st.plotly_chart(fig_bar_seg, use_container_width=True)

                            with st.expander("Ver Registro Detallado de Calificación de Operaciones"):
                                df_seguridad_filtrado = df_bar_data[df_bar_data['TIPO_OPERACION'] == 'POR SEGURIDAD'].copy()
                                if not df_seguridad_filtrado.empty:
                                    st.dataframe(df_seguridad_filtrado, use_container_width=True)
                                else:
                                    st.info("No hay registros de operación POR SEGURIDAD en este periodo.")
                        else:
                            st.warning("Las centrales filtradas no registraron operaciones calificadas en este periodo.")
                    else:
                        st.warning("Faltan datos válidos de INICIO o FIN para graficar las horas de operación.")

                # ==========================================
                # LÓGICA DE EXPORTACIÓN A WORD (EN EL CONTENEDOR INICIAL)
                # ==========================================
                with contenedor_reporte:
                    if st.button("📄 Preparar Reporte Word (Incluye Gráficos)", use_container_width=True):
                        with st.spinner("Compilando gráficos... (Nota: Este proceso requiere tener instalada la librería 'kaleido')"):
                            doc = Document()
                            doc.add_heading('Dashboard de Supervisión - Despacho SEIN', 0)
                            doc.add_paragraph(f"Reporte generado automáticamente el: {datetime.now().strftime('%d/%m/%Y a las %H:%M')}")

                            # SOLUCIÓN: Usamos un diccionario mutable para eludir los problemas de alcance (scope)
                            estado_exportacion = {"error_graficos": False}

                            def agregar_grafico(doc, fig, titulo):
                                if fig is not None:
                                    doc.add_heading(titulo, level=1)
                                    try:
                                        # Se reduce la escala a 1.2 para evitar bloqueos de memoria (Timeout)
                                        img_bytes = fig.to_image(format="png", width=800, height=450, scale=1.2)
                                        imagen_stream = io.BytesIO(img_bytes)
                                        doc.add_picture(imagen_stream, width=Inches(6.0))
                                    except Exception as e:
                                        # Modificamos el valor dentro del diccionario sin usar 'nonlocal'
                                        estado_exportacion["error_graficos"] = True
                                        doc.add_paragraph(f"⚠️ [Error al generar la imagen de este gráfico]")

                            def agregar_tabla(doc, df, titulo):
                                if df is not None and not df.empty:
                                    doc.add_heading(titulo, level=1)
                                    table = doc.add_table(rows=1, cols=len(df.columns))
                                    table.style = 'Table Grid'
                                    for i, col in enumerate(df.columns):
                                        table.rows[0].cells[i].text = str(col)
                                    for _, row in df.iterrows():
                                        row_cells = table.add_row().cells
                                        for i, val in enumerate(row):
                                            row_cells[i].text = str(val)

                            agregar_grafico(doc, fig_tipo, "1. Despacho por Tipo de Generación")
                            agregar_grafico(doc, fig_inter, "2. Flujo de Interconexión")
                            agregar_grafico(doc, fig_cen, "3. Generación SEIN por Central")
                            agregar_grafico(doc, fig_prom, "4. Potencia Promedio Diaria")
                            agregar_grafico(doc, fig_inactividad, "5.1. Horas No Despachadas (Inactividad)")
                            agregar_grafico(doc, fig_actividad, "5.2. Horas de Operación (Diésel/Residual)")
                            agregar_grafico(doc, fig_bar_seg, "6. Calificación de la Operación (Horas por Tipo)")

                            if not df_bar_data.empty:
                                columnas_limpias = df_bar_data.drop(columns=['HORAS_OPERACION', 'CENTRAL_GRUPO'], errors='ignore')
                                agregar_tabla(doc, columnas_limpias, "Registro Detallado de Calificación de Operaciones")

                            buffer_docx = io.BytesIO()
                            doc.save(buffer_docx)
                            buffer_docx.seek(0)
                            
                            # Evaluamos el estado usando nuestro diccionario
                            if estado_exportacion["error_graficos"]:
                                st.error("❌ Ocurrió un error al intentar capturar los gráficos. \n\n**Solución obligatoria:** Abre tu terminal de comandos y ejecuta:\n\n`pip install kaleido==0.1.0.post1`\n\n*(Reinicia tu app después de instalarlo).*")
                            else:
                                st.success("✅ ¡Reporte compilado con éxito incluyendo los gráficos de supervisión!")

                            st.download_button(
                                label="⬇️ Haz clic aquí para descargar tu Reporte (.docx)",
                                data=buffer_docx.getvalue(),
                                file_name=f"Reporte_Despacho_SEIN_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                type="primary",
                                use_container_width=True
                            )
                # ==========================================
                # 8. EVOLUCIÓN Y COMPORTAMIENTO DE LA DEMANDA
                # ==========================================
                st.markdown("---")
                st.header("8. 🌍 Evolución y Comportamiento de la Demanda por Áreas")
                
                if df_dem_raw is None or df_dem_raw.empty:
                    st.info("No se encontraron datos de demanda en el periodo descargado.")
                else:
                    df_subareas = df_dem_raw[df_dem_raw['ÁREA'] != 'SEIN'].copy()
                    df_sein = df_dem_raw[df_dem_raw['ÁREA'] == 'SEIN'].copy()
                    
                    colores_area = {"NORTE": "#FF9900", "CENTRO": "#3366CC", "SUR": "#DC3912"}
                    
                    # Cambio a px.line para generar series de líneas punteadas en lugar de áreas
                    fig_demanda = px.line(
                        df_subareas, x="FECHA_HORA", y="DEMANDA_MW", color="ÁREA",
                        title="Evolución de la Demanda por Áreas Operativas (MW)",
                        color_discrete_map=colores_area,
                        template="plotly_white"
                    )
                    
                    # Configurar las líneas como punteadas (dot)
                    fig_demanda.update_traces(hovertemplate="%{y:,.2f} MW", line=dict(width=2, dash='dot'))
                    
                    # Función para hallar y graficar los puntos Máximos y Mínimos dinámicamente
                    def graficar_min_max(fig, df_filtro, color_marcador, nombre_area):
                        if not df_filtro.empty:
                            idx_max = df_filtro['DEMANDA_MW'].idxmax()
                            idx_min = df_filtro['DEMANDA_MW'].idxmin()
                            
                            max_row = df_filtro.loc[idx_max]
                            min_row = df_filtro.loc[idx_min]
                            
                            # Marcador Máximo (Triángulo apuntando arriba)
                            fig.add_scatter(
                                x=[max_row['FECHA_HORA']], y=[max_row['DEMANDA_MW']],
                                mode='markers+text',
                                marker=dict(color=color_marcador, size=12, symbol='triangle-up'),
                                text=[f"<b>Máx: {max_row['DEMANDA_MW']:,.0f} MW<b>"],
                                textposition="top center",
                                name=f'Máx {nombre_area}',
                                hoverinfo='skip',
                                showlegend=False,
                                textfont=dict(color="blue")
                            )
                            
                            
                    # Iterar para marcar puntos críticos por cada área operativa
                    for area in df_subareas['ÁREA'].unique():
                        graficar_min_max(fig_demanda, df_subareas[df_subareas['ÁREA'] == area], colores_area.get(area, 'blue'), area)
                    
                    if not df_sein.empty:
                        # La línea total del SEIN se mantiene sólida/gruesa o en formato 'dash' para contrastar con los 'dots'
                        fig_demanda.add_scatter(
                            x=df_sein['FECHA_HORA'], y=df_sein['DEMANDA_MW'], mode='lines',
                            line=dict(width=3, color='black', dash='dash'), name='<b>⚡ DEMANDA SEIN TOTAL</b>',
                            hovertemplate='<b>🗓️ %{x|%d/%m/%Y %H:%M} ➡️ %{y:,.2f} MW</b>'
                        )
                        # Marcar máximo y mínimo absoluto del SEIN
                        graficar_min_max(fig_demanda, df_sein, 'blue', 'SEIN')
                    
                    fig_demanda.update_layout(
                        hovermode="x unified",
                        xaxis=dict(tickformat="%d/%m\n%H:%M", title="Fecha Operativa"),
                        yaxis=dict(title="Demanda Activa (MW)"),
                        height=550, margin=dict(t=50, b=50, l=50, r=20)
                    )
                    st.plotly_chart(fig_demanda, use_container_width=True)

                # ==========================================
                # 9. TRAZABILIDAD (DATA CRUDA)
                # ==========================================
                st.markdown("---")
                st.header("9. 🗄️ Trazabilidad de Potencia (Data Cruda - SEIN)")
                
                with st.expander("Ver Matriz de Despacho de Generación", expanded=False):
                    df_pivot = df_plot_cen.copy()
                    df_pivot['FECHA'] = df_pivot['FECHA_HORA'].dt.strftime('%d/%m/%Y')
                    df_pivot['HORA'] = df_pivot['FECHA_HORA'].dt.strftime('%H:%M')
                    
                    jerarquia_columnas = ['CENTRAL'] if es_formato_anexo1 else ['TIPO_CENTRAL', 'EMPRESA', 'CENTRAL']
                    
                    df_matricial = df_pivot.pivot_table(
                        index=['FECHA', 'HORA'],
                        columns=jerarquia_columnas,
                        values='DESPACHO_MW',
                        aggfunc='sum'
                    ).round(2).fillna(0)
                    
                    st.dataframe(df_matricial, use_container_width=True)
                    
                    buffer_xls = io.BytesIO()
                    with pd.ExcelWriter(buffer_xls, engine='openpyxl') as writer:
                        df_matricial.to_excel(writer, sheet_name='Despacho_SEIN')
                        
                    st.download_button(
                        label="📥 Descargar Vista Matricial (Excel)",
                        data=buffer_xls.getvalue(),
                        file_name=f"matriz_sein_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
                # ==========================================
                # 10. ANÁLISIS DE BALANCE DE POTENCIA - ÁREA NORTE
                # ==========================================
                st.markdown("---")
                st.header("10. 📉 Balance de Potencia - Área Norte")
                st.info("Evolución temporal del comportamiento eléctrico en el Norte del país: Superposición de la Demanda del Área Norte (calculada como la suma de los valores absolutos de Generación y Flujo), la Generación local total del Norte y el inverso multiplicativo del Flujo de Interconexión Centro-Norte (-1 * Flujo C-N).")
                
                df_dem_raw = st.session_state.get('df_demanda', pd.DataFrame())
                df_inter_raw = st.session_state.get('df_interconexiones', pd.DataFrame())
                df_despacho_raw = st.session_state.get('df_despacho', pd.DataFrame())
                
                if df_dem_raw.empty or df_inter_raw.empty or df_despacho_raw.empty:
                    st.info("Se requieren datos consolidados de Demanda, Enlaces y Despacho en el periodo para compilar el balance del Área Norte.")
                else:
                    import plotly.graph_objects as go
                    
                    # 1. Agrupar la Demanda real registrada para el Área Norte (Servirá como base de tiempo)
                    df_dem_norte = df_dem_raw[df_dem_raw['ÁREA'] == 'NORTE'].groupby('FECHA_HORA', as_index=False)['DEMANDA_MW'].sum()
                    
                    # 2. Capturar el Flujo Neto Centro-Norte e invertir su signo (* -1)
                    df_cn = df_inter_raw[df_inter_raw['ENLACE'] == 'CENTRO-NORTE'].copy()
                    df_cn_total = df_cn.groupby('FECHA_HORA', as_index=False)['FLUJO_MW'].sum()
                    df_cn_total['FLUJO_NEG'] = df_cn_total['FLUJO_MW'] * -1
                    
                    # 3. Agrupar la Generación total local en la Zona Norte (usando la data maestra limpia)
                    df_gen_norte = df_despacho_raw[df_despacho_raw['ZONA'] == 'NORTE'].groupby('FECHA_HORA', as_index=False)['DESPACHO_MW'].sum()
                    
                    # Consolidar las tres series de tiempo en una única matriz temporal
                    df_balance = df_dem_norte.merge(df_cn_total[['FECHA_HORA', 'FLUJO_NEG']], on='FECHA_HORA', how='inner')
                    df_balance = df_balance.merge(df_gen_norte, on='FECHA_HORA', how='inner')
                    df_balance.columns = ['FECHA_HORA', 'DEMANDA', 'FLUJO_NEG', 'GENERACION']
                    
                    # --- CÁLCULO CORREGIDO DE LA DEMANDA ---
                    # Demanda Norte = |Generación Norte| + |-1 * Flujo C-N|
                    df_balance['DEMANDA'] = df_balance['GENERACION'].abs() + df_balance['FLUJO_NEG'].abs()
                    
                    fig_balance = go.Figure()
                    
                    # Curva 1: Demanda Área Norte (Línea Sólida Gruesa)
                    fig_balance.add_trace(go.Scatter(
                        x=df_balance['FECHA_HORA'], y=df_balance['DEMANDA'],
                        mode='lines', line=dict(width=3, color='#FF9900'),
                        name='<b>📉 DEMANDA NORTE</b>',
                        hovertemplate="<b>Demanda Norte</b>: %{y:,.2f} MW"
                    ))
                    
                    # Curva 2: Generación Local Norte (Línea Punteada Gruesa)
                    fig_balance.add_trace(go.Scatter(
                        x=df_balance['FECHA_HORA'], y=df_balance['GENERACION'],
                        mode='lines', line=dict(width=3, dash='dot', color='#1f77b4'),
                        name='<b>🏭 GENERACIÓN NORTE</b>',
                        hovertemplate="<b>Generación Norte</b>: %{y:,.2f} MW"
                    ))
                    
                    # Curva 3: Flujo Inverso Centro-Norte (Línea Discontinua Gruesa)
                    fig_balance.add_trace(go.Scatter(
                        x=df_balance['FECHA_HORA'], y=df_balance['FLUJO_NEG'],
                        mode='lines', line=dict(width=3, dash='dash', color='#9467bd'),
                        name='<b>🔌 -1 * FLUJO C-N</b>',
                        hovertemplate="<b>-1 * Flujo C-N</b>: %{y:,.2f} MW"
                    ))
                    
                    # Subrutina para inyectar marcadores de extremos (Máximos/Mínimos) en negrita, azules y con 2 decimales
                    def marcar_extremos_balance(fig, x_data, y_data, color_marcador):
                        if not y_data.empty:
                            idx_max = y_data.idxmax()
                            idx_min = y_data.idxmin()
                            
                            fig.add_trace(go.Scatter(
                                x=[x_data[idx_max]], y=[y_data[idx_max]],
                                mode='markers+text', marker=dict(color=color_marcador, size=10, symbol='triangle-up'),
                                text=[f"<b>Máx: {y_data[idx_max]:,.2f}</b>"], textposition="top center", # <--- 2 DECIMALES
                                showlegend=False, hoverinfo='skip',
                                textfont=dict(color="blue")
                            ))
                            fig.add_trace(go.Scatter(
                                x=[x_data[idx_min]], y=[y_data[idx_min]],
                                mode='markers+text', marker=dict(color=color_marcador, size=10, symbol='triangle-down'),
                                text=[f"<b>Mín: {y_data[idx_min]:,.2f}</b>"], textposition="bottom center", # <--- 2 DECIMALES
                                showlegend=False, hoverinfo='skip',
                                textfont=dict(color="blue")
                            ))
                    
                    # Marcar picos y valles para cada una de las series
                    marcar_extremos_balance(fig_balance, df_balance['FECHA_HORA'], df_balance['DEMANDA'], '#FF9900')
                    marcar_extremos_balance(fig_balance, df_balance['FECHA_HORA'], df_balance['GENERACION'], '#1f77b4')
                    marcar_extremos_balance(fig_balance, df_balance['FECHA_HORA'], df_balance['FLUJO_NEG'], '#9467bd')
                    
                    # Definición dinámica de límites de visualización con +400 MW de holgura
                    max_absoluto = max(df_balance['DEMANDA'].max(), df_balance['GENERACION'].max(), df_balance['FLUJO_NEG'].max())
                    min_absoluto = min(df_balance['DEMANDA'].min(), df_balance['GENERACION'].min(), df_balance['FLUJO_NEG'].min())
                    limite_y_sup = max_absoluto + 400
                    limite_y_inf = min_absoluto * 1.15 if min_absoluto < 0 else 0
                    
                    fig_balance.update_layout(
                        hovermode="x unified",
                        height=600, 
                        margin=dict(t=50, b=50, l=50, r=150), # Margen derecho amplio para la leyenda vertical
                        legend=dict(
                            title="<b>Variables del Área</b>",
                            orientation="v",
                            yanchor="top", 
                            y=1, 
                            xanchor="left", 
                            x=1.05
                        ),
                        template="plotly_white",
                        xaxis=dict(title_text="Fecha Operativa", tickformat="%d/%m\n%H:%M"),
                        yaxis=dict(title_text="Potencia Activa (MW)", range=[limite_y_inf, limite_y_sup])
                    )
                    
                    st.plotly_chart(fig_balance, use_container_width=True)
                    
                    # Tabla matricial complementaria para trazabilidad de datos
                    with st.expander("Ver Datos de Balance Norte (Vista Matricial)"):
                        df_mat_b = df_balance.copy()
                        df_mat_b['FECHA'] = df_mat_b['FECHA_HORA'].dt.strftime('%d/%m/%Y')
                        df_mat_b['HORA'] = df_mat_b['FECHA_HORA'].dt.strftime('%H:%M')
                        matriz_b = df_mat_b.pivot_table(
                            index=['FECHA', 'HORA'], 
                            values=['DEMANDA', 'GENERACION', 'FLUJO_NEG'], 
                            aggfunc='mean'
                        ).round(2)
                        # Reordenar columnas para legibilidad
                        matriz_b = matriz_b[['DEMANDA', 'GENERACION', 'FLUJO_NEG']]
                        st.dataframe(matriz_b, use_container_width=True)

                # ==========================================
                # 11. EVOLUCIÓN GLOBAL DE COSTOS MARGINALES (CMg)
                # ==========================================
                st.markdown("---")
                st.header("11. 📈 Evolución Consolidada de Costos Marginales")
                st.info("Comparativa directa del Costo Marginal (USD/MWh) en las principales barras de referencia del Norte, Centro y Sur del SEIN.")
                
                df_cmg_raw = st.session_state.get('df_cmg', pd.DataFrame())
                
                if df_cmg_raw.empty:
                    st.info("No se encontraron datos de Costos Marginales en el periodo descargado.")
                else:
                    df_cmg_plot_11 = df_cmg_raw.sort_values(['FECHA_HORA', 'BARRA']).copy()
                    
                    # Paleta de colores estándar para las 3 barras
                    colores_barra_11 = {
                        'SANTA ROSA 220': '#d62728',  # Rojo (Centro)
                        'MONTALVO 220': '#2ca02c',    # Verde (Sur)
                        'TRUJILLO 220': '#ff7f0e'     # Naranja (Norte)
                    }
                    
                    # Generación de la gráfica de líneas múltiples
                    fig_cmg_11 = px.line(
                        df_cmg_plot_11, x="FECHA_HORA", y="CMG_USD", color="BARRA",
                        title="Costo Marginal Consolidado por Barra (USD/MWh)",
                        labels={"BARRA": "Barra 220 kV", "CMG_USD": "Costo Marginal (USD/MWh)"},
                        color_discrete_map=colores_barra_11, 
                        template="plotly_white"
                    )
                    
                    # Aplicar formato de línea gruesa y punteada
                    fig_cmg_11.update_traces(hovertemplate="<b>%{data.name}</b>: %{y:,.2f} USD/MWh", line=dict(width=3, dash='dot'))
                    
                    # Subrutina para añadir marcadores de extremos con texto AZUL
                    def graficar_extremos_azules(fig, df_filtro, color_marcador, nombre_barra):
                        if not df_filtro.empty:
                            idx_max = df_filtro['CMG_USD'].idxmax()
                            idx_min = df_filtro['CMG_USD'].idxmin()
                            
                            # Marcador Máximo (Texto Azul)
                            fig.add_scatter(
                                x=[df_filtro.loc[idx_max, 'FECHA_HORA']], y=[df_filtro.loc[idx_max, 'CMG_USD']],
                                mode='markers+text', marker=dict(color=color_marcador, size=12, symbol='triangle-up'),
                                text=[f"<b>Máx: {df_filtro.loc[idx_max, 'CMG_USD']:,.1f}</b>"], 
                                textposition="top center",
                                textfont=dict(color="blue"), # <--- ETIQUETA EN AZUL
                                name=f'Máx {nombre_barra}', hoverinfo='skip', showlegend=False
                            )
                            
                            # Marcador Mínimo (Texto Azul)
                            fig.add_scatter(
                                x=[df_filtro.loc[idx_min, 'FECHA_HORA']], y=[df_filtro.loc[idx_min, 'CMG_USD']],
                                mode='markers+text', marker=dict(color=color_marcador, size=12, symbol='triangle-down'),
                                text=[f"<b>Mín: {df_filtro.loc[idx_min, 'CMG_USD']:,.1f}</b>"], 
                                textposition="bottom center",
                                textfont=dict(color="blue"), # <--- ETIQUETA EN AZUL
                                name=f'Mín {nombre_barra}', hoverinfo='skip', showlegend=False
                            )

                    # Iterar sobre todas las barras disponibles en la data original
                    for barra in df_cmg_plot_11['BARRA'].unique():
                        df_barra_11 = df_cmg_plot_11[df_cmg_plot_11['BARRA'] == barra]
                        graficar_extremos_azules(fig_cmg_11, df_barra_11, colores_barra_11.get(barra, 'black'), barra)
                    
                    # Ajuste de escala dinámica con margen visual del 15%
                    max_val_cmg_11 = df_cmg_plot_11['CMG_USD'].max()
                    limite_superior_11 = max_val_cmg_11 * 1.15 if max_val_cmg_11 > 0 else 50
                    
                    fig_cmg_11.update_layout(
                        hovermode="x unified",
                        xaxis=dict(tickformat="%d/%m\n%H:%M", title="Fecha Operativa"),
                        yaxis=dict(title="Costo Marginal (USD/MWh)", range=[0, limite_superior_11]),
                        height=600, 
                        margin=dict(t=50, b=50, l=50, r=150),
                        legend=dict(
                            title="<b>Barras SEIN</b>",
                            orientation="v", 
                            yanchor="top", 
                            y=1, 
                            xanchor="left", 
                            x=1.02
                        )
                    )
                    
                    st.plotly_chart(fig_cmg_11, use_container_width=True)
                    
                    with st.expander("Ver Datos Consolidados de CMg (Vista Matricial)"):
                        df_cmg_pivot_11 = df_cmg_plot_11.copy()
                        df_cmg_pivot_11['FECHA'] = df_cmg_pivot_11['FECHA_HORA'].dt.strftime('%d/%m/%Y')
                        df_cmg_pivot_11['HORA'] = df_cmg_pivot_11['FECHA_HORA'].dt.strftime('%H:%M')
                        matriz_cmg_11 = df_cmg_pivot_11.pivot_table(index=['FECHA', 'HORA'], columns=['BARRA'], values='CMG_USD', aggfunc='mean').round(2)
                        st.dataframe(matriz_cmg_11, use_container_width=True)